"""
File content extraction service.
Extracts text from various file formats including PDF, Word, Excel, images, and text files.
"""

import io
import os
from typing import Optional, Tuple
import base64

# PDF extraction
try:
    from pypdf import PdfReader
    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

# Word document extraction
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# Excel extraction
try:
    import pandas as pd
    HAS_PANDAS = True
except ImportError:
    HAS_PANDAS = False

# Image OCR (optional - requires tesseract installed on system)
try:
    from PIL import Image
    import pytesseract
    HAS_OCR = True
except ImportError:
    HAS_OCR = False


# Maximum content length to store (characters)
MAX_CONTENT_LENGTH = 100000


def extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from a PDF file."""
    if not HAS_PYPDF:
        return "[PDF extraction not available - pypdf not installed]"
    
    try:
        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)
        
        text_parts = []
        for page_num, page in enumerate(reader.pages, 1):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"--- Page {page_num} ---\n{page_text}")
        
        return "\n\n".join(text_parts) if text_parts else "[No text content found in PDF]"
    except Exception as e:
        return f"[Error extracting PDF: {str(e)}]"


def extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from a Word document (.docx)."""
    if not HAS_DOCX:
        return "[DOCX extraction not available - python-docx not installed]"
    
    try:
        doc_file = io.BytesIO(file_content)
        doc = DocxDocument(doc_file)
        
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract tables
        for table_idx, table in enumerate(doc.tables, 1):
            table_text = [f"\n--- Table {table_idx} ---"]
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_text.append(" | ".join(row_data))
            text_parts.append("\n".join(table_text))
        
        return "\n\n".join(text_parts) if text_parts else "[No text content found in document]"
    except Exception as e:
        return f"[Error extracting DOCX: {str(e)}]"


def extract_text_from_excel(file_content: bytes, filename: str) -> str:
    """Extract text from Excel files (.xlsx, .xls, .csv)."""
    if not HAS_PANDAS:
        return "[Excel extraction not available - pandas not installed]"
    
    try:
        file_obj = io.BytesIO(file_content)
        
        # Determine file type and read accordingly
        if filename.endswith('.csv'):
            df = pd.read_csv(file_obj)
            sheets = {'CSV Data': df}
        else:
            # Excel file - read all sheets
            excel_file = pd.ExcelFile(file_obj)
            sheets = {sheet: pd.read_excel(excel_file, sheet_name=sheet) 
                     for sheet in excel_file.sheet_names}
        
        text_parts = []
        for sheet_name, df in sheets.items():
            text_parts.append(f"=== Sheet: {sheet_name} ===")
            
            # Convert DataFrame to string representation
            # Include column headers and data
            text_parts.append(df.to_string(index=False, max_rows=500))
            
            # Add summary stats for numeric columns
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                text_parts.append("\n--- Summary Statistics ---")
                text_parts.append(df[numeric_cols].describe().to_string())
        
        return "\n\n".join(text_parts) if text_parts else "[No data found in spreadsheet]"
    except Exception as e:
        return f"[Error extracting Excel: {str(e)}]"


def extract_text_from_image(file_content: bytes) -> str:
    """Extract text from images using OCR."""
    if not HAS_OCR:
        return "[Image OCR not available - pytesseract or PIL not installed]"
    
    try:
        image = Image.open(io.BytesIO(file_content))
        
        # Perform OCR
        text = pytesseract.image_to_string(image)
        
        if text.strip():
            return f"[Text extracted from image via OCR]\n\n{text}"
        else:
            return "[No text found in image - image may not contain readable text]"
    except pytesseract.TesseractNotFoundError:
        return "[OCR not available - Tesseract is not installed on the system]"
    except Exception as e:
        return f"[Error extracting text from image: {str(e)}]"


def extract_text_from_text_file(file_content: bytes) -> str:
    """Extract text from plain text files."""
    try:
        # Try UTF-8 first, then fall back to other encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252', 'ascii']:
            try:
                return file_content.decode(encoding)
            except UnicodeDecodeError:
                continue
        
        # Last resort - decode with errors ignored
        return file_content.decode('utf-8', errors='ignore')
    except Exception as e:
        return f"[Error reading text file: {str(e)}]"


def extract_content_from_file(file_content: bytes, filename: str, content_type: str) -> Tuple[str, str]:
    """
    Extract text content from a file based on its type.
    
    Returns:
        Tuple of (extracted_text, detected_file_category)
    """
    filename_lower = filename.lower()
    
    # Determine file type and extract accordingly
    if filename_lower.endswith('.pdf') or content_type == 'application/pdf':
        text = extract_text_from_pdf(file_content)
        category = 'pdf'
    
    elif filename_lower.endswith('.docx') or content_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
        text = extract_text_from_docx(file_content)
        category = 'word'
    
    elif filename_lower.endswith('.doc') or content_type == 'application/msword':
        # Old .doc format - limited support
        text = "[Legacy .doc format - please convert to .docx for full text extraction]"
        category = 'word'
    
    elif filename_lower.endswith(('.xlsx', '.xls')) or content_type in [
        'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        'application/vnd.ms-excel'
    ]:
        text = extract_text_from_excel(file_content, filename)
        category = 'spreadsheet'
    
    elif filename_lower.endswith('.csv') or content_type == 'text/csv':
        text = extract_text_from_excel(file_content, filename)
        category = 'spreadsheet'
    
    elif filename_lower.endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.tiff', '.webp')) or content_type.startswith('image/'):
        text = extract_text_from_image(file_content)
        category = 'image'
    
    elif filename_lower.endswith(('.txt', '.md', '.json', '.xml', '.html', '.css', '.js', '.py')) or content_type.startswith('text/'):
        text = extract_text_from_text_file(file_content)
        category = 'text'
    
    elif filename_lower.endswith(('.pptx', '.ppt')):
        # PowerPoint - basic support
        text = "[PowerPoint files - please export as PDF for text extraction]"
        category = 'presentation'
    
    else:
        # Try to read as text, fall back to binary info
        try:
            text = file_content.decode('utf-8')
            category = 'text'
        except UnicodeDecodeError:
            text = f"[Binary file - {len(file_content)} bytes, type: {content_type}]"
            category = 'binary'
    
    # Truncate if too long
    if len(text) > MAX_CONTENT_LENGTH:
        text = text[:MAX_CONTENT_LENGTH] + f"\n\n[Content truncated - showing first {MAX_CONTENT_LENGTH} characters of {len(text)} total]"
    
    return text, category


def get_supported_extensions() -> list:
    """Return list of supported file extensions."""
    extensions = [
        '.pdf',      # PDF documents
        '.docx',     # Word documents
        '.doc',      # Legacy Word (limited)
        '.xlsx',     # Excel spreadsheets
        '.xls',      # Legacy Excel
        '.csv',      # CSV files
        '.txt',      # Text files
        '.md',       # Markdown
        '.json',     # JSON
        '.xml',      # XML
        '.html',     # HTML
        '.png',      # Images (OCR)
        '.jpg',
        '.jpeg',
        '.gif',
        '.bmp',
        '.tiff',
        '.webp',
        '.pptx',     # PowerPoint (limited)
        '.ppt',
    ]
    return extensions

