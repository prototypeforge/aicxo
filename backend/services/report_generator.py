"""
Report generation service for creating PDF and DOCX exports of C-Suite meeting reports.
Supports two styles: colorful (matching the app UI with dark background) and professional (formal report style).
"""
import io
import markdown2
from datetime import datetime
from typing import Dict, List, Any, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak, BaseDocTemplate, PageTemplate, Frame
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
import os


# Color schemes for DOCX
COLORFUL_SCHEME = {
    "primary": RGBColor(248, 196, 113),  # Gold
    "secondary": RGBColor(89, 130, 245),  # Sapphire
    "accent": RGBColor(147, 112, 219),  # Purple
    "success": RGBColor(74, 222, 128),  # Green
    "warning": RGBColor(251, 191, 36),  # Amber
    "danger": RGBColor(239, 68, 68),  # Red
    "text": RGBColor(229, 231, 235),  # Light gray text
    "text_bright": RGBColor(255, 255, 255),  # White
    "muted": RGBColor(156, 163, 175),  # Gray
    "background": "18181B",  # Dark background (hex without #)
}

PROFESSIONAL_SCHEME = {
    "primary": RGBColor(31, 41, 55),  # Dark gray
    "secondary": RGBColor(55, 65, 81),  # Medium gray
    "accent": RGBColor(37, 99, 235),  # Blue
    "success": RGBColor(5, 150, 105),  # Green
    "warning": RGBColor(217, 119, 6),  # Orange
    "danger": RGBColor(220, 38, 38),  # Red
    "text": RGBColor(17, 24, 39),  # Near black
    "text_bright": RGBColor(17, 24, 39),  # Near black
    "muted": RGBColor(107, 114, 128),  # Gray
    "background": None,  # White (default)
}

# PDF Color schemes
COLORFUL_PDF = {
    "primary": colors.HexColor("#FFB420"),  # Gold (brighter)
    "secondary": colors.HexColor("#5982F5"),  # Sapphire
    "accent": colors.HexColor("#A78BFA"),  # Purple (brighter)
    "success": colors.HexColor("#4ADE80"),  # Green
    "warning": colors.HexColor("#FBBF24"),  # Amber
    "danger": colors.HexColor("#EF4444"),  # Red
    "text": colors.HexColor("#E5E7EB"),  # Light gray text
    "text_bright": colors.HexColor("#FFFFFF"),  # White
    "muted": colors.HexColor("#9CA3AF"),  # Gray
    "background": colors.HexColor("#0F0F23"),  # Dark obsidian background
    "card_bg": colors.HexColor("#1F1F35"),  # Card background
}

PROFESSIONAL_PDF = {
    "primary": colors.HexColor("#1F2937"),  # Dark gray
    "secondary": colors.HexColor("#374151"),  # Medium gray
    "accent": colors.HexColor("#2563EB"),  # Blue
    "success": colors.HexColor("#059669"),  # Green
    "warning": colors.HexColor("#D97706"),  # Orange
    "danger": colors.HexColor("#DC2626"),  # Red
    "text": colors.HexColor("#111827"),  # Near black
    "text_bright": colors.HexColor("#111827"),  # Near black
    "muted": colors.HexColor("#6B7280"),  # Gray
    "background": colors.HexColor("#FFFFFF"),  # White
    "card_bg": colors.HexColor("#F9FAFB"),  # Near white
}


def strip_markdown(text: str) -> str:
    """Convert markdown to plain text for simple formatting."""
    if not text:
        return ""
    # Convert markdown to HTML, then strip tags
    html = markdown2.markdown(text)
    # Simple tag stripping
    import re
    clean = re.compile('<.*?>')
    return re.sub(clean, '', html).strip()


def format_confidence(confidence: float) -> str:
    """Format confidence as percentage."""
    return f"{int(confidence * 100)}%"


def get_confidence_color(confidence: float, style: str) -> RGBColor:
    """Get color based on confidence level."""
    scheme = COLORFUL_SCHEME if style == "colorful" else PROFESSIONAL_SCHEME
    if confidence >= 0.8:
        return scheme["success"]
    elif confidence >= 0.6:
        return scheme["warning"]
    return RGBColor(239, 68, 68)  # Red


_fonts_registered = False
_font_names = ('Helvetica', 'Helvetica-Bold', 'Helvetica-Oblique')


def _register_fonts():
    """Register Unicode-compatible fonts for PDF generation."""
    global _fonts_registered, _font_names
    
    # Only register once
    if _fonts_registered:
        return _font_names
    
    # Try to find and register fonts with Unicode support
    font_configs = [
        # DejaVu fonts (best Unicode support, common on Linux)
        {
            'regular': [
                '/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf',
                '/usr/share/fonts/dejavu/DejaVuSans.ttf',
                'C:/Windows/Fonts/DejaVuSans.ttf',
            ],
            'bold': [
                '/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf',
                '/usr/share/fonts/dejavu/DejaVuSans-Bold.ttf',
                'C:/Windows/Fonts/DejaVuSans-Bold.ttf',
            ],
            'italic': [
                '/usr/share/fonts/truetype/dejavu/DejaVuSans-Oblique.ttf',
                '/usr/share/fonts/dejavu/DejaVuSans-Oblique.ttf',
                'C:/Windows/Fonts/DejaVuSans-Oblique.ttf',
            ],
            'names': ('DejaVuSans', 'DejaVuSans-Bold', 'DejaVuSans-Oblique'),
        },
        # Liberation Sans (common on Linux, similar to Arial)
        {
            'regular': [
                '/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf',
                '/usr/share/fonts/liberation/LiberationSans-Regular.ttf',
            ],
            'bold': [
                '/usr/share/fonts/truetype/liberation/LiberationSans-Bold.ttf',
                '/usr/share/fonts/liberation/LiberationSans-Bold.ttf',
            ],
            'italic': [
                '/usr/share/fonts/truetype/liberation/LiberationSans-Italic.ttf',
                '/usr/share/fonts/liberation/LiberationSans-Italic.ttf',
            ],
            'names': ('LiberationSans', 'LiberationSans-Bold', 'LiberationSans-Italic'),
        },
        # FreeSans (common fallback on Linux)
        {
            'regular': [
                '/usr/share/fonts/truetype/freefont/FreeSans.ttf',
                '/usr/share/fonts/freefont/FreeSans.ttf',
            ],
            'bold': [
                '/usr/share/fonts/truetype/freefont/FreeSansBold.ttf',
                '/usr/share/fonts/freefont/FreeSansBold.ttf',
            ],
            'italic': [
                '/usr/share/fonts/truetype/freefont/FreeSansOblique.ttf',
                '/usr/share/fonts/freefont/FreeSansOblique.ttf',
            ],
            'names': ('FreeSans', 'FreeSans-Bold', 'FreeSans-Italic'),
        },
        # Arial (Windows/macOS)
        {
            'regular': [
                'C:/Windows/Fonts/arial.ttf',
                '/Library/Fonts/Arial.ttf',
                '/System/Library/Fonts/Supplemental/Arial.ttf',
            ],
            'bold': [
                'C:/Windows/Fonts/arialbd.ttf',
                '/Library/Fonts/Arial Bold.ttf',
            ],
            'italic': [
                'C:/Windows/Fonts/ariali.ttf',
                '/Library/Fonts/Arial Italic.ttf',
            ],
            'names': ('ArialUnicode', 'ArialUnicode-Bold', 'ArialUnicode-Italic'),
        },
    ]
    
    for config in font_configs:
        try:
            regular_path = None
            bold_path = None
            italic_path = None
            
            # Find regular font
            for path in config['regular']:
                if os.path.exists(path):
                    regular_path = path
                    break
            
            if not regular_path:
                continue
            
            # Find bold font
            for path in config['bold']:
                if os.path.exists(path):
                    bold_path = path
                    break
            
            # Find italic font
            for path in config['italic']:
                if os.path.exists(path):
                    italic_path = path
                    break
            
            # Register fonts
            names = config['names']
            pdfmetrics.registerFont(TTFont(names[0], regular_path))
            
            if bold_path:
                pdfmetrics.registerFont(TTFont(names[1], bold_path))
            else:
                # Use regular as fallback for bold
                pdfmetrics.registerFont(TTFont(names[1], regular_path))
            
            if italic_path:
                pdfmetrics.registerFont(TTFont(names[2], italic_path))
            else:
                # Use regular as fallback for italic
                pdfmetrics.registerFont(TTFont(names[2], regular_path))
            
            _fonts_registered = True
            _font_names = names
            return names
            
        except Exception as e:
            # Log but continue trying other fonts
            print(f"Failed to register font {config['names'][0]}: {e}")
            continue
    
    # Return Helvetica as last resort (limited Unicode support)
    _fonts_registered = True
    _font_names = ('Helvetica', 'Helvetica-Bold', 'Helvetica-Oblique')
    return _font_names


def _draw_background(canvas, doc, bg_color):
    """Draw background color on the page."""
    canvas.saveState()
    canvas.setFillColor(bg_color)
    canvas.rect(0, 0, doc.pagesize[0], doc.pagesize[1], fill=True, stroke=False)
    canvas.restoreState()


async def generate_pdf_report(
    meeting: Dict[str, Any],
    style: str = "colorful"
) -> bytes:
    """Generate a PDF report for the meeting."""
    buffer = io.BytesIO()
    
    colors_scheme = COLORFUL_PDF if style == "colorful" else PROFESSIONAL_PDF
    is_colorful = style == "colorful"
    
    # Register Unicode-compatible fonts
    font_regular, font_bold, font_italic = _register_fonts()
    
    # Create PDF document with custom page template for background
    doc = BaseDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=0.75*inch,
        leftMargin=0.75*inch,
        topMargin=0.75*inch,
        bottomMargin=0.75*inch
    )
    
    # Create frame for content
    frame = Frame(
        doc.leftMargin,
        doc.bottomMargin,
        doc.width,
        doc.height,
        id='normal'
    )
    
    # Create page template with background
    def on_page(canvas, doc):
        if is_colorful:
            _draw_background(canvas, doc, colors_scheme["background"])
    
    template = PageTemplate(id='main', frames=frame, onPage=on_page)
    doc.addPageTemplates([template])
    
    # Define styles with Unicode-compatible fonts
    styles = getSampleStyleSheet()
    
    # Title style
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=20,
        alignment=TA_CENTER,
        textColor=colors_scheme["primary"],
        fontName=font_bold
    )
    
    # Heading styles
    h1_style = ParagraphStyle(
        'CustomH1',
        parent=styles['Heading1'],
        fontSize=18,
        spaceAfter=12,
        spaceBefore=20,
        textColor=colors_scheme["primary"],
        fontName=font_bold
    )
    
    h2_style = ParagraphStyle(
        'CustomH2',
        parent=styles['Heading2'],
        fontSize=14,
        spaceAfter=8,
        spaceBefore=15,
        textColor=colors_scheme["secondary"],
        fontName=font_bold
    )
    
    h3_style = ParagraphStyle(
        'CustomH3',
        parent=styles['Heading3'],
        fontSize=12,
        spaceAfter=6,
        spaceBefore=12,
        textColor=colors_scheme["accent"],
        fontName=font_bold
    )
    
    # Body text - use light text for colorful, dark for professional
    body_style = ParagraphStyle(
        'CustomBody',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=8,
        alignment=TA_JUSTIFY,
        textColor=colors_scheme["text"],
        fontName=font_regular,
        leading=16
    )
    
    # Muted text
    muted_style = ParagraphStyle(
        'CustomMuted',
        parent=styles['Normal'],
        fontSize=10,
        textColor=colors_scheme["muted"],
        fontName=font_regular,
        spaceAfter=6
    )
    
    # Quote/highlight style - gold for colorful, dark for professional
    quote_style = ParagraphStyle(
        'CustomQuote',
        parent=styles['Normal'],
        fontSize=11,
        leftIndent=20,
        borderPadding=10,
        textColor=colors_scheme["primary"] if is_colorful else colors_scheme["text"],
        fontName=font_italic,
        spaceAfter=12,
        spaceBefore=8
    )
    
    # Bright text style for important content
    bright_style = ParagraphStyle(
        'CustomBright',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=8,
        alignment=TA_JUSTIFY,
        textColor=colors_scheme["text_bright"],
        fontName=font_regular,
        leading=16
    )
    
    # Build content
    story = []
    
    # Title
    story.append(Paragraph("CxO Ninja Meeting Report", title_style))
    story.append(Spacer(1, 0.2*inch))
    
    # Meeting metadata
    created_at = meeting.get('created_at', '')
    if isinstance(created_at, str):
        try:
            created_at = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
        except:
            created_at = datetime.now()
    
    completed_at = meeting.get('completed_at')
    if completed_at and isinstance(completed_at, str):
        try:
            completed_at = datetime.fromisoformat(completed_at.replace('Z', '+00:00'))
        except:
            completed_at = None
    
    meta_text = f"Created: {created_at.strftime('%B %d, %Y at %I:%M %p') if created_at else 'N/A'}"
    if completed_at:
        meta_text += f" | Completed: {completed_at.strftime('%B %d, %Y at %I:%M %p')}"
    if meeting.get('current_version', 1) > 1:
        meta_text += f" | Version {meeting.get('current_version')}"
    
    story.append(Paragraph(meta_text, muted_style))
    story.append(Spacer(1, 0.3*inch))
    
    # Question Section
    story.append(Paragraph("Question Presented", h1_style))
    story.append(Paragraph(meeting.get('question', 'No question provided'), body_style))
    
    if meeting.get('context'):
        story.append(Spacer(1, 0.1*inch))
        story.append(Paragraph("<b>Additional Context:</b>", muted_style))
        story.append(Paragraph(meeting.get('context'), body_style))
    
    story.append(Spacer(1, 0.3*inch))
    
    # Chair's Summary Section
    story.append(Paragraph("Chair of the Board's Summary", h1_style))
    
    if meeting.get('chair_summary'):
        story.append(Paragraph("<b>Board Discussion Summary</b>", h3_style))
        summary_text = strip_markdown(meeting.get('chair_summary', ''))
        # Split by paragraphs
        for para in summary_text.split('\n\n'):
            if para.strip():
                story.append(Paragraph(para.strip(), body_style))
    
    if meeting.get('chair_recommendation'):
        story.append(Spacer(1, 0.15*inch))
        story.append(Paragraph("<b>Official Recommendation</b>", h3_style))
        rec_text = strip_markdown(meeting.get('chair_recommendation', ''))
        for para in rec_text.split('\n\n'):
            if para.strip():
                story.append(Paragraph(para.strip(), quote_style))
    
    story.append(Spacer(1, 0.3*inch))
    
    # Board Member Opinions
    opinions = meeting.get('opinions', [])
    if opinions:
        story.append(Paragraph(f"Board Member Opinions ({len(opinions)})", h1_style))
        
        for opinion in opinions:
            # Agent header
            agent_name = opinion.get('agent_name', 'Unknown Agent')
            agent_role = opinion.get('agent_role', 'Board Member')
            confidence = opinion.get('confidence', 0)
            
            story.append(Paragraph(f"<b>{agent_name}</b> - {agent_role}", h2_style))
            story.append(Paragraph(f"Confidence: {format_confidence(confidence)}", muted_style))
            
            # Opinion
            story.append(Paragraph("<b>Opinion:</b>", h3_style))
            opinion_text = strip_markdown(opinion.get('opinion', ''))
            for para in opinion_text.split('\n\n'):
                if para.strip():
                    story.append(Paragraph(para.strip(), body_style))
            
            # Reasoning
            story.append(Paragraph("<b>Reasoning:</b>", h3_style))
            reasoning_text = strip_markdown(opinion.get('reasoning', ''))
            for para in reasoning_text.split('\n\n'):
                if para.strip():
                    story.append(Paragraph(para.strip(), body_style))
            
            # Expertise
            weights = opinion.get('weights_applied', {})
            high_weights = [(k.replace('_', ' ').title(), v) for k, v in weights.items() if v > 0.3]
            if high_weights:
                high_weights.sort(key=lambda x: x[1], reverse=True)
                expertise_str = ", ".join([f"{k}: {int(v*100)}%" for k, v in high_weights])
                story.append(Paragraph(f"<i>Key Expertise: {expertise_str}</i>", muted_style))
            
            story.append(Spacer(1, 0.2*inch))
    
    # Follow-up Questions
    follow_ups = meeting.get('follow_ups', [])
    if follow_ups:
        story.append(Paragraph(f"Follow-up Questions ({len(follow_ups)})", h1_style))
        
        for i, fu in enumerate(follow_ups, 1):
            story.append(Paragraph(f"<b>Q{i}:</b> {fu.get('question', '')}", h3_style))
            if fu.get('chair_response'):
                story.append(Paragraph("<b>Chair's Response:</b>", muted_style))
                response_text = strip_markdown(fu.get('chair_response', ''))
                for para in response_text.split('\n\n'):
                    if para.strip():
                        story.append(Paragraph(para.strip(), body_style))
            story.append(Spacer(1, 0.15*inch))
    
    # Footer
    story.append(Spacer(1, 0.5*inch))
    story.append(Paragraph(
        f"<i>Report generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}</i>",
        muted_style
    ))
    
    # Build PDF
    doc.build(story)
    
    buffer.seek(0)
    return buffer.getvalue()


def set_document_background(doc, color_hex: str):
    """Set the background color of the entire Word document."""
    # Access the document's settings element
    settings = doc.settings.element
    
    # Create background element
    bg = OxmlElement('w:background')
    bg.set(qn('w:color'), color_hex)
    
    # Find displayBackgroundShape or create it
    display_bg = settings.find(qn('w:displayBackgroundShape'))
    if display_bg is None:
        display_bg = OxmlElement('w:displayBackgroundShape')
        settings.insert(0, display_bg)
    
    # Insert background at the beginning of settings
    settings.insert(0, bg)


def add_colored_paragraph(doc, text: str, color: RGBColor, bold: bool = False, italic: bool = False, size: int = 11):
    """Add a paragraph with specific text color."""
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.color.rgb = color
    run.font.bold = bold
    run.font.italic = italic
    run.font.size = Pt(size)
    return para


async def generate_docx_report(
    meeting: Dict[str, Any],
    style: str = "colorful"
) -> bytes:
    """Generate a DOCX report for the meeting."""
    doc = Document()
    
    colors_scheme = COLORFUL_SCHEME if style == "colorful" else PROFESSIONAL_SCHEME
    is_colorful = style == "colorful"
    
    # Set dark background for colorful style
    if is_colorful and colors_scheme.get("background"):
        set_document_background(doc, colors_scheme["background"])
    
    # Set document properties
    core_props = doc.core_properties
    core_props.title = "CxO Ninja Meeting Report"
    core_props.author = "CxO Ninja - Your Digital C-Suite"
    
    # Title
    title = doc.add_heading("CxO Ninja Meeting Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = colors_scheme["primary"]
        run.font.size = Pt(28)
    
    # Metadata
    created_at = meeting.get('created_at', '')
    if isinstance(created_at, str):
        try:
            created_at = datetime.fromisoformat(created_at.replace('Z', '+00:00'))
        except:
            created_at = datetime.now()
    
    completed_at = meeting.get('completed_at')
    if completed_at and isinstance(completed_at, str):
        try:
            completed_at = datetime.fromisoformat(completed_at.replace('Z', '+00:00'))
        except:
            completed_at = None
    
    meta_para = doc.add_paragraph()
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_text = f"Created: {created_at.strftime('%B %d, %Y at %I:%M %p') if created_at else 'N/A'}"
    if completed_at:
        meta_text += f" | Completed: {completed_at.strftime('%B %d, %Y at %I:%M %p')}"
    if meeting.get('current_version', 1) > 1:
        meta_text += f" | Version {meeting.get('current_version')}"
    meta_run = meta_para.add_run(meta_text)
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = colors_scheme["muted"]
    
    doc.add_paragraph()
    
    # Question Section
    question_heading = doc.add_heading("Question Presented", 1)
    for run in question_heading.runs:
        run.font.color.rgb = colors_scheme["primary"]
    
    q_para = doc.add_paragraph()
    q_run = q_para.add_run(meeting.get('question', 'No question provided'))
    q_run.font.color.rgb = colors_scheme["text"]
    q_para.paragraph_format.space_after = Pt(12)
    
    if meeting.get('context'):
        context_label = doc.add_paragraph()
        label_run = context_label.add_run("Additional Context:")
        label_run.font.bold = True
        label_run.font.color.rgb = colors_scheme["muted"]
        
        ctx_para = doc.add_paragraph()
        ctx_run = ctx_para.add_run(meeting.get('context'))
        ctx_run.font.color.rgb = colors_scheme["text"]
    
    # Chair's Summary
    chair_heading = doc.add_heading("Chair of the Board's Summary", 1)
    for run in chair_heading.runs:
        run.font.color.rgb = colors_scheme["primary"]
    
    if meeting.get('chair_summary'):
        summary_label = doc.add_heading("Board Discussion Summary", 2)
        for run in summary_label.runs:
            run.font.color.rgb = colors_scheme["secondary"]
        
        summary_text = strip_markdown(meeting.get('chair_summary', ''))
        for para in summary_text.split('\n\n'):
            if para.strip():
                p = doc.add_paragraph()
                r = p.add_run(para.strip())
                r.font.color.rgb = colors_scheme["text"]
    
    if meeting.get('chair_recommendation'):
        rec_label = doc.add_heading("Official Recommendation", 2)
        for run in rec_label.runs:
            run.font.color.rgb = colors_scheme["primary"] if is_colorful else colors_scheme["accent"]
        
        rec_text = strip_markdown(meeting.get('chair_recommendation', ''))
        for para in rec_text.split('\n\n'):
            if para.strip():
                rec_para = doc.add_paragraph()
                rec_run = rec_para.add_run(para.strip())
                rec_run.font.italic = True
                rec_run.font.color.rgb = colors_scheme["primary"] if is_colorful else colors_scheme["text"]
    
    # Board Member Opinions
    opinions = meeting.get('opinions', [])
    if opinions:
        opinions_heading = doc.add_heading(f"Board Member Opinions ({len(opinions)})", 1)
        for run in opinions_heading.runs:
            run.font.color.rgb = colors_scheme["primary"]
        
        for opinion in opinions:
            agent_name = opinion.get('agent_name', 'Unknown Agent')
            agent_role = opinion.get('agent_role', 'Board Member')
            confidence = opinion.get('confidence', 0)
            
            # Agent header
            agent_heading = doc.add_heading(f"{agent_name} - {agent_role}", 2)
            for run in agent_heading.runs:
                run.font.color.rgb = colors_scheme["secondary"]
            
            # Confidence
            conf_para = doc.add_paragraph()
            conf_run = conf_para.add_run(f"Confidence: {format_confidence(confidence)}")
            conf_run.font.size = Pt(10)
            conf_run.font.color.rgb = get_confidence_color(confidence, style)
            
            # Opinion
            opinion_label = doc.add_paragraph()
            opinion_label_run = opinion_label.add_run("Opinion:")
            opinion_label_run.font.bold = True
            opinion_label_run.font.color.rgb = colors_scheme["text_bright"]
            
            opinion_text = strip_markdown(opinion.get('opinion', ''))
            for para in opinion_text.split('\n\n'):
                if para.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(para.strip())
                    r.font.color.rgb = colors_scheme["text"]
            
            # Reasoning
            reasoning_label = doc.add_paragraph()
            reasoning_label_run = reasoning_label.add_run("Reasoning:")
            reasoning_label_run.font.bold = True
            reasoning_label_run.font.color.rgb = colors_scheme["text_bright"]
            
            reasoning_text = strip_markdown(opinion.get('reasoning', ''))
            for para in reasoning_text.split('\n\n'):
                if para.strip():
                    p = doc.add_paragraph()
                    r = p.add_run(para.strip())
                    r.font.color.rgb = colors_scheme["text"]
            
            # Expertise
            weights = opinion.get('weights_applied', {})
            high_weights = [(k.replace('_', ' ').title(), v) for k, v in weights.items() if v > 0.3]
            if high_weights:
                high_weights.sort(key=lambda x: x[1], reverse=True)
                expertise_str = ", ".join([f"{k}: {int(v*100)}%" for k, v in high_weights])
                expertise_para = doc.add_paragraph()
                expertise_run = expertise_para.add_run(f"Key Expertise: {expertise_str}")
                expertise_run.font.italic = True
                expertise_run.font.size = Pt(10)
                expertise_run.font.color.rgb = colors_scheme["accent"] if is_colorful else colors_scheme["muted"]
            
            doc.add_paragraph()
    
    # Follow-up Questions
    follow_ups = meeting.get('follow_ups', [])
    if follow_ups:
        fu_heading = doc.add_heading(f"Follow-up Questions ({len(follow_ups)})", 1)
        for run in fu_heading.runs:
            run.font.color.rgb = colors_scheme["primary"]
        
        for i, fu in enumerate(follow_ups, 1):
            q_para = doc.add_paragraph()
            q_label = q_para.add_run(f"Q{i}: ")
            q_label.font.bold = True
            q_label.font.color.rgb = colors_scheme["text_bright"]
            q_text = q_para.add_run(fu.get('question', ''))
            q_text.font.color.rgb = colors_scheme["text"]
            
            if fu.get('chair_response'):
                response_label = doc.add_paragraph()
                response_label_run = response_label.add_run("Chair's Response:")
                response_label_run.font.bold = True
                response_label_run.font.color.rgb = colors_scheme["primary"]
                
                response_text = strip_markdown(fu.get('chair_response', ''))
                for para in response_text.split('\n\n'):
                    if para.strip():
                        p = doc.add_paragraph()
                        r = p.add_run(para.strip())
                        r.font.color.rgb = colors_scheme["text"]
            
            doc.add_paragraph()
    
    # Footer
    doc.add_paragraph()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run(
        f"Report generated on {datetime.now().strftime('%B %d, %Y at %I:%M %p')}"
    )
    footer_run.font.italic = True
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = colors_scheme["muted"]
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

